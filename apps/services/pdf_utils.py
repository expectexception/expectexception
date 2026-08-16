"""
PDF conversion utilities — multi-engine strategy for maximum accuracy.

Conversion strategy:
  PDF → DOCX  pdf2docx (best layout fidelity: tables, images, columns)
              → fallback: convert_pdf_to_docx_native() (PyMuPDF + python-docx)
  PDF → other soffice (LibreOffice), isolated user profile per call
  DOC/DOCX/ODT/RTF/TXT → PDF  soffice with fidelity-tuned export options
  OCR (scanned)  page-at-a-time render → tesseract → reassembled searchable PDF

Two invariants worth knowing before editing:

  soffice exits 0 even when an export fails, so never trust its return code —
  check that the output file exists.

  soffice loads a PDF into Draw, not Writer, so it cannot export a PDF through
  any Writer filter. PDF→DOCX via LibreOffice is impossible by construction,
  which is why the fallback is a native rebuild rather than another soffice
  call.
"""

import io
import logging
import os
import shutil
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import Any

from django.conf import settings

logger = logging.getLogger(__name__)


class PDFConversionError(Exception):
    """Raised when PDF conversion fails."""

    pass


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def get_soffice_path() -> str:
    """Return path to LibreOffice soffice binary."""
    if hasattr(settings, "SOFFICE_CMD") and settings.SOFFICE_CMD:
        if os.path.exists(settings.SOFFICE_CMD):
            return settings.SOFFICE_CMD

    for path in ["/usr/bin/soffice", "/usr/bin/libreoffice"]:
        if os.path.exists(path):
            return path

    raise PDFConversionError(
        "LibreOffice not found. Install with: sudo apt-get install libreoffice"
    )


def validate_pdf_file(file_path: str, max_size: int | None = None) -> None:
    """Validate PDF file exists, is a PDF, and is within size limits."""
    if not os.path.exists(file_path):
        raise PDFConversionError(f"Input file not found: {file_path}")

    if not file_path.lower().endswith(".pdf"):
        raise PDFConversionError("Input file must have a .pdf extension")

    if max_size is None:
        max_size = getattr(settings, "PDF_MAX_FILE_SIZE", 50 * 1024 * 1024)

    file_size = os.path.getsize(file_path)
    if file_size > max_size:
        raise PDFConversionError(
            f"File too large ({file_size / (1024*1024):.1f} MB). "
            f"Maximum: {max_size / (1024*1024):.0f} MB"
        )


def _soffice_env(home_dir: str) -> dict:
    """Build a clean environment for soffice with an isolated user profile."""
    env = {
        "PATH": "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin",
        # Each call gets its own HOME so LibreOffice profiles never clash
        "HOME": home_dir,
        # Suppress D-Bus / display requirements
        "DISPLAY": "",
        "DBUS_SESSION_BUS_ADDRESS": "disabled:",
    }
    for key in ("LANG", "LC_ALL", "LC_CTYPE", "TZ"):
        if key in os.environ:
            env[key] = os.environ[key]
    return env


# ---------------------------------------------------------------------------
# soffice conversion
# ---------------------------------------------------------------------------

# Correct LibreOffice export filter names
_SOFFICE_FILTERS = {
    "docx": "MS Word 2007 XML",
    "doc": "MS Word 97",
    "odt": "writer8",
    "rtf": "Rich Text Format",
    "txt": "Text (encoded)",
    "pdf": "writer_pdf_Export",
}

_SOFFICE_EXTENSIONS = {
    "docx": ".docx",
    "doc": ".doc",
    "odt": ".odt",
    "rtf": ".rtf",
    "txt": ".txt",
    "pdf": ".pdf",
}


def convert_pdf_with_soffice(
    input_pdf: str,
    output_format: str,
    output_path: str | None = None,
    timeout: int = 180,
) -> str:
    """
    Convert PDF → another format using LibreOffice (soffice).

    Uses an isolated HOME directory per call to avoid LibreOffice
    user-profile corruption when multiple conversions run in parallel.

    Returns: path to the converted file.
    """
    validate_pdf_file(input_pdf)

    output_format = output_format.lower()
    if output_format not in _SOFFICE_FILTERS:
        raise PDFConversionError(
            f"Unsupported format '{output_format}'. " f"Supported: {', '.join(_SOFFICE_FILTERS)}"
        )

    extension = _SOFFICE_EXTENSIONS[output_format]
    filter_name = _SOFFICE_FILTERS[output_format]
    soffice_path = get_soffice_path()

    # Create isolated temp dirs for this conversion
    run_dir = tempfile.mkdtemp(prefix="soffice_run_")
    try:
        # soffice writes output as <stem><ext> in --outdir
        out_dir = run_dir
        expected_stem = Path(input_pdf).stem
        expected_output = os.path.join(out_dir, f"{expected_stem}{extension}")

        cmd = [
            soffice_path,
            "--headless",
            "--norestore",
            "--nofirststartwizard",
            "--convert-to",
            f"{output_format}:{filter_name}",
            "--outdir",
            out_dir,
            input_pdf,
        ]

        logger.info(f"soffice: {input_pdf} → {output_format}  (filter: {filter_name})")
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
            env=_soffice_env(run_dir),
        )

        # soffice exits 0 even when the export fails (verified: converting a PDF
        # with a Writer filter prints "Error: Please verify input parameters"
        # and still returns 0), so the return code alone proves nothing — the
        # real test is whether an output file appeared.
        if result.returncode != 0:
            msg = (result.stderr or result.stdout or "no output").strip()
            raise PDFConversionError(f"LibreOffice failed (rc={result.returncode}): {msg}")

        if not os.path.exists(expected_output):
            # soffice sometimes keeps input stem 1:1 — do a glob search
            candidates = list(Path(out_dir).glob(f"*{extension}"))
            if not candidates:
                raise PDFConversionError(
                    f"soffice produced no output file in {out_dir}. "
                    f"stdout: {result.stdout}  stderr: {result.stderr}"
                )
            expected_output = str(candidates[0])

        # Move to the caller-specified output path (or MEDIA_ROOT/converted)
        if output_path is None:
            converted_dir = os.path.join(settings.MEDIA_ROOT, "converted")
            os.makedirs(converted_dir, exist_ok=True)
            output_path = os.path.join(converted_dir, Path(expected_output).name)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        shutil.move(expected_output, output_path)

        logger.info(f"soffice conversion done: {output_path}")
        return output_path

    except subprocess.TimeoutExpired:
        raise PDFConversionError(f"Conversion timed out after {timeout}s. Try a smaller file.")
    finally:
        shutil.rmtree(run_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Document → PDF conversion
# ---------------------------------------------------------------------------

# LibreOffice's default PDF export downsamples images to 300 DPI, JPEG-encodes
# them, and does not guarantee font embedding — all of which lose fidelity
# against the source document. These options turn that off and additionally
# emit a tagged PDF, which carries the heading/table structure through, so a
# later PDF→DOCX round trip has something to rebuild from.
# Measured on a 2400x1600 source image: the defaults downsampled it to
# 1950x1300 and re-encoded it as lossy JPEG. Lossless keeps line art, diagrams
# and screenshots exact — and for such content the PDF is usually *smaller*
# too (1290 KB -> 101 KB in that test). Photographic documents compress worse
# losslessly, so it stays configurable.
_PDF_EXPORT_OPTIONS = {
    "UseLosslessCompression": os.getenv("PDF_EXPORT_LOSSLESS", "True") == "True",
    "ReduceImageResolution": False,
    "EmbedStandardFonts": True,
    "UseTaggedPDF": True,
    "ExportBookmarks": True,
    "ExportNotes": False,
}

# Input filters that LibreOffice guesses badly when left to itself.
_SOFFICE_INFILTERS = {
    # Without this the encoding is sniffed from the bytes and non-ASCII text
    # comes out as mojibake.
    ".txt": "Text (encoded):UTF8",
}

_DOC_TO_PDF_EXTENSIONS = (".doc", ".docx", ".odt", ".rtf", ".txt")


def _pdf_export_filter() -> str:
    """Build the `pdf:writer_pdf_Export:{json}` argument for --convert-to."""
    import json

    options = {
        key: {"type": "boolean", "value": str(value).lower()}
        for key, value in _PDF_EXPORT_OPTIONS.items()
    }
    return f"pdf:writer_pdf_Export:{json.dumps(options)}"


def convert_document_to_pdf(
    input_doc: str,
    output_path: str,
    timeout: int = 180,
) -> str:
    """Convert DOC/DOCX/ODT/RTF/TXT → PDF via LibreOffice.

    Kept here rather than inline in the view so both conversion directions
    share one set of soffice invariants (resolved binary path, isolated HOME,
    D-Bus/display suppressed, output-file existence checked rather than the
    exit code, which soffice reports as 0 even on a failed export).
    """
    if not os.path.exists(input_doc):
        raise PDFConversionError(f"Input file not found: {input_doc}")

    extension = Path(input_doc).suffix.lower()
    if extension not in _DOC_TO_PDF_EXTENSIONS:
        raise PDFConversionError(
            f"Unsupported format '{extension}'. Supported: {', '.join(_DOC_TO_PDF_EXTENSIONS)}"
        )

    soffice_path = get_soffice_path()
    run_dir = tempfile.mkdtemp(prefix="doc2pdf_run_")
    try:
        cmd = [soffice_path, "--headless", "--norestore", "--nofirststartwizard"]
        infilter = _SOFFICE_INFILTERS.get(extension)
        if infilter:
            cmd.append(f"--infilter={infilter}")
        cmd += ["--convert-to", _pdf_export_filter(), "--outdir", run_dir, input_doc]

        logger.info("soffice: %s → pdf", input_doc)
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=timeout, env=_soffice_env(run_dir)
        )

        produced = os.path.join(run_dir, f"{Path(input_doc).stem}.pdf")
        if not os.path.exists(produced):
            candidates = list(Path(run_dir).glob("*.pdf"))
            if not candidates:
                msg = (result.stderr or result.stdout or "no output").strip()
                raise PDFConversionError(f"LibreOffice produced no PDF: {msg}")
            produced = str(candidates[0])

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        shutil.move(produced, output_path)
        logger.info("doc→pdf conversion done: %s", output_path)
        return output_path

    except subprocess.TimeoutExpired:
        raise PDFConversionError(f"Conversion timed out after {timeout}s. Try a smaller file.")
    finally:
        shutil.rmtree(run_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# pdf2docx conversion
# ---------------------------------------------------------------------------


def convert_pdf_with_pdf2docx(
    input_pdf: str,
    output_docx: str,
    password: str | None = None,
    start_page: int = 0,
    end_page: int | None = None,
) -> str:
    """
    Convert PDF → DOCX using pdf2docx (best layout/table fidelity).

    Falls back details:
      - Preserves text, tables, images, columns, headers/footers
      - Handles multi-column layouts better than soffice

    Returns: path to output DOCX.
    """
    validate_pdf_file(input_pdf)
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)

    try:
        from pdf2docx import Converter
    except ImportError:
        raise PDFConversionError("pdf2docx not installed. Run: pip install pdf2docx")

    logger.info(f"pdf2docx: {input_pdf} → {output_docx}")
    cv = None
    try:
        cv = Converter(input_pdf, password=password)
        cv.convert(output_docx, start=start_page, end=end_page)
        cv.close()
    except Exception as exc:
        if cv:
            try:
                cv.close()
            except Exception:
                pass
        raise PDFConversionError(f"pdf2docx error: {exc}")

    if not os.path.exists(output_docx):
        raise PDFConversionError("pdf2docx did not produce an output file")

    logger.info(f"pdf2docx done: {output_docx}")
    return output_docx


# ---------------------------------------------------------------------------
# Native PDF → DOCX rebuild (fallback when pdf2docx cannot parse a file)
# ---------------------------------------------------------------------------

# Bit flags PyMuPDF sets on each text span.
_FITZ_FLAG_ITALIC = 1 << 1
_FITZ_FLAG_BOLD = 1 << 4


def _dominant_font_size(page) -> float:
    """Most-used span size on the page — the body-text size headings stand out from."""
    from collections import Counter

    sizes = Counter()
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        for line in block["lines"]:
            for span in line["spans"]:
                if span["text"].strip():
                    sizes[round(span["size"], 1)] += len(span["text"])
    return sizes.most_common(1)[0][0] if sizes else 11.0


def convert_pdf_to_docx_native(input_pdf: str, output_docx: str) -> str:
    """Rebuild a PDF as a real DOCX using PyMuPDF for extraction.

    This exists because the previous fallback — asking LibreOffice for
    'docx:MS Word 2007 XML' — cannot work by construction: LibreOffice opens a
    PDF in Draw, and a Draw document cannot be saved through a Writer filter.
    It logged "Error: Please verify input parameters", wrote nothing, and still
    exited 0. So whenever pdf2docx failed, the "fallback" failed too and the
    whole conversion errored out.

    Produces flowing, editable text (not one text box per line): headings from
    relative font size, bold/italic runs, and real Word tables.
    """
    import fitz
    from docx import Document
    from docx.shared import Pt

    validate_pdf_file(input_pdf)
    os.makedirs(os.path.dirname(output_docx) or ".", exist_ok=True)

    doc = fitz.open(input_pdf)
    out = Document()

    try:
        for page_index, page in enumerate(doc):
            if page_index > 0:
                out.add_page_break()

            body_size = _dominant_font_size(page)

            # Tables are emitted separately, so remember their regions and skip
            # any text inside them — otherwise every cell is duplicated as a
            # loose paragraph next to the table.
            table_rects = []
            try:
                for table in page.find_tables():
                    rows = table.extract()
                    if not rows or not any(any(c for c in r) for r in rows):
                        continue
                    table_rects.append(fitz.Rect(table.bbox))
                    docx_table = out.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                    docx_table.style = "Table Grid"
                    for r, row in enumerate(rows):
                        for c, cell in enumerate(row):
                            if c < len(docx_table.columns):
                                docx_table.cell(r, c).text = (cell or "").strip()
            except Exception as e:
                logger.warning("Table extraction failed on page %d: %s", page_index, e)

            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                block_rect = fitz.Rect(block["bbox"])
                if any(block_rect.intersects(tr) for tr in table_rects):
                    continue

                for line in block["lines"]:
                    spans = [s for s in line["spans"] if s["text"].strip()]
                    if not spans:
                        continue

                    largest = max(s["size"] for s in spans)
                    if largest >= body_size * 1.6:
                        paragraph = out.add_heading("", level=1)
                    elif largest >= body_size * 1.2:
                        paragraph = out.add_heading("", level=2)
                    else:
                        paragraph = out.add_paragraph()

                    for span in spans:
                        run = paragraph.add_run(span["text"])
                        run.bold = bool(span["flags"] & _FITZ_FLAG_BOLD)
                        run.italic = bool(span["flags"] & _FITZ_FLAG_ITALIC)
                        run.font.size = Pt(round(span["size"], 1))

        out.save(output_docx)
    finally:
        doc.close()

    if not os.path.exists(output_docx):
        raise PDFConversionError("Native PDF→DOCX rebuild produced no output file")

    logger.info("native pdf→docx rebuild done: %s", output_docx)
    return output_docx


# ---------------------------------------------------------------------------
# OCR pipeline (pytesseract → reconstructed PDF → convert)
# ---------------------------------------------------------------------------


def perform_ocr_on_pdf(
    input_pdf: str,
    output_pdf: str,
    language: str = "eng",
    dpi: int = 300,
    timeout: int = 300,
) -> str:
    """
    OCR a scanned PDF into a searchable PDF using PyMuPDF + pytesseract.

    Renders, OCRs and appends one page at a time. The previous implementation
    called pdf2image.convert_from_path() for the whole document, which
    materializes every page as a PIL image before any OCR starts — at 300 DPI
    that is roughly 25 MB of RAM per A4 page, so a 100-page scan needed
    gigabytes and would OOM the worker. It also called
    image_to_pdf_or_hocr() twice per page and threw the first result away,
    doubling the tesseract time for every document.

    Returns: path to OCR'd PDF.
    """
    validate_pdf_file(input_pdf)
    os.makedirs(os.path.dirname(output_pdf) or ".", exist_ok=True)

    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise PDFConversionError(f"OCR dependencies missing: {e}")

    logger.info("OCR: %s (lang=%s, dpi=%d)", input_pdf, language, dpi)

    source = fitz.open(input_pdf)
    result = fitz.open()
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)

    try:
        for _index, page in enumerate(source):
            pixmap = page.get_pixmap(matrix=matrix)
            with Image.open(io.BytesIO(pixmap.tobytes("png"))) as image:
                pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                    image,
                    lang=language,
                    extension="pdf",
                    config="--psm 3",  # Fully automatic page segmentation
                )
            del pixmap

            with fitz.open("pdf", pdf_bytes) as page_pdf:
                result.insert_pdf(page_pdf)

        if result.page_count == 0:
            raise PDFConversionError("OCR produced no pages")

        result.save(output_pdf)
        logger.info("OCR complete: %s (%d pages)", output_pdf, result.page_count)
        return output_pdf

    except PDFConversionError:
        raise
    except Exception as exc:
        raise PDFConversionError(f"OCR pipeline error: {exc}")
    finally:
        result.close()
        source.close()


# ---------------------------------------------------------------------------
# High-level smart converter (used by the Celery task)
# ---------------------------------------------------------------------------


def smart_convert_pdf(
    input_pdf: str,
    output_format: str,
    output_path: str,
    ocr_enabled: bool = False,
    ocr_lang: str = "eng",
) -> dict[str, Any]:
    """
    Smart PDF conversion with multi-engine fallback strategy.

    Strategy:
      1. If ocr_enabled: run pytesseract OCR pipeline first → searchable PDF
      2. For DOCX: try pdf2docx first (best quality), fallback to soffice
      3. For other formats: soffice directly

    Returns dict with: file_path, format, converted_size, ocr_used, engine_used
    """
    validate_pdf_file(input_pdf)
    original_size = os.path.getsize(input_pdf)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    working_pdf = input_pdf
    ocr_actually_used = False

    # Step 1: OCR if requested
    if ocr_enabled:
        # Was output_path.replace(f'.{output_format}', '_ocr_tmp.pdf'), which
        # replaces *every* occurrence of the format string in the path and does
        # nothing at all when the suffix case differs (".DOCX"). When it matched
        # nothing the OCR result was written straight to output_path, so the
        # converter then read and wrote the same file.
        ocr_output = str(Path(output_path).with_suffix("")) + "_ocr_tmp.pdf"
        try:
            working_pdf = perform_ocr_on_pdf(input_pdf, ocr_output, language=ocr_lang)
            ocr_actually_used = True
            logger.info("OCR step completed successfully")
        except PDFConversionError as e:
            logger.warning(f"OCR step failed, converting without OCR: {e}")
            working_pdf = input_pdf  # fallback to original

    # Step 2: Convert to target format
    engine_used = None

    if output_format.lower() == "docx":
        # Try pdf2docx first (better layout preservation), then rebuild the
        # document natively. The old fallback here asked LibreOffice for a
        # Writer filter on a PDF, which it loads in Draw — that combination
        # cannot produce a file, so a pdf2docx failure used to fail the request.
        try:
            final_path = convert_pdf_with_pdf2docx(working_pdf, output_path)
            engine_used = "pdf2docx"
        except PDFConversionError as e:
            logger.warning(f"pdf2docx failed ({e}), rebuilding natively")
            final_path = convert_pdf_to_docx_native(working_pdf, output_path)
            engine_used = "native-fallback"
    else:
        final_path = convert_pdf_with_soffice(working_pdf, output_format, output_path)
        engine_used = "soffice"

    # Clean up OCR temp
    if ocr_actually_used and os.path.exists(working_pdf) and working_pdf != input_pdf:
        try:
            os.remove(working_pdf)
        except OSError:
            pass

    converted_size = os.path.getsize(final_path)

    return {
        "file_path": final_path,
        "format": output_format.upper(),
        "original_size": original_size,
        "converted_size": converted_size,
        "ocr_used": ocr_actually_used,
        "engine_used": engine_used,
    }


# ---------------------------------------------------------------------------
# Batch helper (unchanged API)
# ---------------------------------------------------------------------------


def batch_convert_pdf(
    input_pdf: str,
    output_formats: list,
    output_dir: str | None = None,
) -> dict[str, dict[str, Any]]:
    """Convert PDF to multiple formats in one call."""
    validate_pdf_file(input_pdf)

    if output_dir is None:
        output_dir = os.path.join(settings.MEDIA_ROOT, "converted")
    os.makedirs(output_dir, exist_ok=True)

    results = {}
    for fmt in output_formats:
        try:
            output_path = os.path.join(
                output_dir, f"{Path(input_pdf).stem}_{uuid.uuid4().hex[:6]}.{fmt.lower()}"
            )
            info = smart_convert_pdf(input_pdf, fmt, output_path)
            results[fmt] = {
                "status": "success",
                "path": info["file_path"],
                "size": info["converted_size"],
                "url": f"{settings.MEDIA_URL}converted/{os.path.basename(info['file_path'])}",
                "engine": info["engine_used"],
            }
        except PDFConversionError as exc:
            results[fmt] = {"status": "failed", "error": str(exc)}

    return results
